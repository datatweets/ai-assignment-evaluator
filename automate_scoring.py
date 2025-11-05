#!/usr/bin/env python3
"""
Automated Assignment Scoring System
This script evaluates student assignments using AI and generates structured feedback.
"""

import os
import json
import anthropic
from pathlib import Path
from docx import Document
from datetime import datetime
from typing import Dict, Any, List


class AssignmentScorer:
    """Handles the automated scoring of student assignments."""

    def __init__(self, config_path: str = "config.json"):
        """Initialize the scorer with configuration."""
        self.config = self._load_config(config_path)
        self.evaluation_prompt = self._load_evaluation_prompt()
        self.client = anthropic.Anthropic(api_key=self.config.get("anthropic_api_key"))

    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration from JSON file."""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"Warning: {config_path} not found. Using environment variables.")
            return {
                "anthropic_api_key": os.getenv("ANTHROPIC_API_KEY", ""),
                "model": "claude-3-5-sonnet-20241022",
                "max_tokens": 4096
            }

    def _load_evaluation_prompt(self) -> Dict[str, Any]:
        """Load the evaluation prompt from JSON file."""
        with open("evaluation_prompt.json", 'r', encoding='utf-8') as f:
            return json.load(f)

    def _build_evaluation_prompt(self, student_text: str) -> str:
        """Build the complete evaluation prompt with student submission."""
        prompt_parts = [
            self.evaluation_prompt["system_role"],
            "",
            "The student completed an assignment with this structure:",
            ""
        ]

        # Add assignment structure
        for section_key, section_data in self.evaluation_prompt["assignment_structure"].items():
            prompt_parts.append(f"- {section_data['title']}")
            if "questions" in section_data:
                for q in section_data["questions"]:
                    prompt_parts.append(f"  - {q}")
            if "steps" in section_data:
                for step in section_data["steps"]:
                    prompt_parts.append(f"  - {step}")
            prompt_parts.append("")

        # Add important note
        prompt_parts.extend([
            "IMPORTANT:",
            self.evaluation_prompt["important_note"],
            "",
            "Your job is to:",
            "1. Read the full student submission.",
            "2. Score it from 0 to 25 using the criteria below.",
            "3. Explain WHY they got that score (for example: 10/25, 15/25, 25/25).",
            "4. Give one concrete suggestion for improvement.",
            "",
            "--------------------------------",
            "EVALUATION CRITERIA (5 points each)",
            "--------------------------------",
            ""
        ])

        # Add evaluation criteria
        for criterion_name, criterion_data in self.evaluation_prompt["evaluation_criteria"].items():
            title = criterion_name.replace("_", " ").title()
            prompt_parts.append(f"{len(prompt_parts) - 15}) {title.upper()} (0–{criterion_data['max_points']})")
            prompt_parts.append("Focus on:")
            for focus_item in criterion_data["focus"]:
                prompt_parts.append(f"- {focus_item}")
            prompt_parts.append("")
            prompt_parts.append("Scoring:")
            for score_range, description in criterion_data["scoring_guide"].items():
                prompt_parts.append(f"- {score_range}: {description}")
            prompt_parts.append("")

        # Add scoring summary
        prompt_parts.extend([
            "-----------------",
            "SCORING SUMMARY",
            "-----------------",
            "",
            "Add the points from the 5 categories (max 25):"
        ])
        for score_range, label in self.evaluation_prompt["scoring_summary"]["ranges"].items():
            prompt_parts.append(f"- {score_range} = {label}")

        prompt_parts.extend([
            "",
            "=" * 60,
            "REQUIRED OUTPUT FORMAT (FOLLOW EXACTLY)",
            "=" * 60,
            "",
            "Your response MUST use this EXACT format with proper spacing and colons.",
            "Do NOT add extra asterisks or change the colon placement.",
            "",
            "EXAMPLE (copy this structure):",
            "",
            "**Total Score:** 15 / 25",
            "",
            "**Category Breakdown:**",
            "- Prompt Clarity: 3/5",
            "- Real-World Relevance: 4/5",
            "- Reflection Quality: 2/5",
            "- Responsible Use: 3/5",
            "- Writing Clarity: 3/5",
            "",
            "**Summary Feedback:**",
            "Your 2-4 sentence feedback here...",
            "",
            "**Improvement Suggestion:**",
            "Your 1-2 sentence suggestion here...",
            "",
            "=" * 60,
            "IMPORTANT FORMAT RULES:",
            "=" * 60,
            "1. Use '**Total Score:**' (with colon after 'Score', no extra asterisks at end)",
            "2. Use '- Category Name: X/5' (space after colon, no extra spaces)",
            "3. Keep section headers exactly as shown above",
            "4. " + self.evaluation_prompt["evaluation_instruction"],
            "",
            "CRITICAL VERIFICATION STEP:",
            "Before scoring, verify the submission contains ACTUAL STUDENT WORK:",
            "- Does it have Section 1 with Q1 and Q2 answers?",
            "- Does it have Section 2 with Steps 1-5 completed by the student?",
            "- Does it have Section 3 with Q1, Q2, Q3 answers?",
            "",
            "If the submission is MISSING these sections (e.g., only contains evaluation tables, headers, or templates),",
            "assign 0-1 points for ALL categories and note that major sections are missing.",
            "",
            "================================",
            "STUDENT SUBMISSION TO EVALUATE:",
            "================================",
            "",
            student_text
        ])

        return "\n".join(prompt_parts)

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text content from a Word document."""
        doc = Document(docx_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)

    def check_if_submission_is_complete(self, student_text: str) -> tuple[bool, str]:
        """Check if the submission contains actual student work or is just a template."""
        text_lower = student_text.lower()

        # Check for required sections
        has_section_1 = "section 1" in text_lower
        has_section_2 = "section 2" in text_lower
        has_section_3 = "section 3" in text_lower

        # Check for key content markers that indicate ACTUAL student answers
        has_q1_answer = ("q1" in text_lower or "q1." in text_lower) and ("generative ai" in text_lower or "artificial intelligence" in text_lower)
        has_step_content = ("step 1" in text_lower and "step 2" in text_lower and ("prompt" in text_lower or "write" in text_lower))

        # Red flags for templates - evaluation table without actual answers
        has_eval_table = ("criterion" in text_lower and "evaluator comments" in text_lower)

        # Count how many section headers vs actual content we have
        section_count = sum([has_section_1, has_section_2, has_section_3])

        # If it has evaluation table but NO section headers at all, it's just a template
        if has_eval_table and section_count == 0:
            return False, "This appears to be an evaluation template with no actual student submission content."

        # Check if we have substantive content (more than just headers and tables)
        content_length = len(student_text)
        if content_length < 300:  # Very short submission
            return False, "Submission is too brief (less than 300 characters) and appears incomplete."

        # If we have section headers but no actual answers to questions
        if (has_section_1 or has_section_2 or has_section_3) and not (has_q1_answer or has_step_content):
            # Check if it's ONLY headers and evaluation tables
            lines = [line.strip() for line in student_text.split('\n') if line.strip()]
            substantive_lines = [line for line in lines if len(line) > 50 and not line.startswith('Section') and 'criterion' not in line.lower()]

            if len(substantive_lines) < 5:  # Less than 5 substantive lines suggests it's mostly empty
                return False, "Submission contains section headers but lacks actual student responses."

        return True, "Submission appears complete"

    def parse_evaluation_response(self, response_text: str) -> Dict[str, Any]:
        """Parse the AI's evaluation response into structured JSON."""
        result = {
            "total_score": "",
            "category_breakdown": {
                "prompt_clarity": "",
                "real_world_relevance": "",
                "reflection_quality": "",
                "responsible_use": "",
                "writing_clarity": ""
            },
            "summary_feedback": "",
            "improvement_suggestion": "",
            "raw_response": response_text
        }

        lines = response_text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()

            if line.startswith("**Total Score"):
                # Handle both "**Total Score:** X / 25" and "**Total Score: X / 25**"
                score_text = line.replace("**Total Score:**", "").replace("**Total Score:", "").replace("**", "").strip()
                result["total_score"] = score_text

            elif line.startswith("- Prompt Clarity:"):
                result["category_breakdown"]["prompt_clarity"] = line.split(":")[-1].strip()

            elif line.startswith("- Real-World Relevance:"):
                result["category_breakdown"]["real_world_relevance"] = line.split(":")[-1].strip()

            elif line.startswith("- Reflection Quality:"):
                result["category_breakdown"]["reflection_quality"] = line.split(":")[-1].strip()

            elif line.startswith("- Responsible Use:"):
                result["category_breakdown"]["responsible_use"] = line.split(":")[-1].strip()

            elif line.startswith("- Writing Clarity:"):
                result["category_breakdown"]["writing_clarity"] = line.split(":")[-1].strip()

            elif line.startswith("**Summary Feedback"):
                current_section = "summary"

            elif line.startswith("**Improvement Suggestion"):
                current_section = "improvement"

            elif current_section == "summary" and line and not line.startswith("**"):
                result["summary_feedback"] += line + " "

            elif current_section == "improvement" and line and not line.startswith("**"):
                result["improvement_suggestion"] += line + " "

        result["summary_feedback"] = result["summary_feedback"].strip()
        result["improvement_suggestion"] = result["improvement_suggestion"].strip()

        # Validate and fix total score if missing
        if not result["total_score"]:
            # Try to calculate from breakdown
            scores = []
            for category_score in result["category_breakdown"].values():
                if category_score and "/" in category_score:
                    try:
                        score_num = int(category_score.split("/")[0].strip())
                        scores.append(score_num)
                    except:
                        pass

            if len(scores) == 5:  # All 5 categories found
                total = sum(scores)
                result["total_score"] = f"{total} / 25"
                print(f"  ⚠️  Total score was missing, calculated from breakdown: {result['total_score']}")

        return result

    def evaluate_assignment(self, student_text: str) -> Dict[str, Any]:
        """Send student assignment to AI for evaluation."""
        prompt = self._build_evaluation_prompt(student_text)

        message = self.client.messages.create(
            model=self.config.get("model", "claude-3-5-sonnet-20241022"),
            max_tokens=self.config.get("max_tokens", 4096),
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        response_text = message.content[0].text
        return self.parse_evaluation_response(response_text)

    def score_assignment(self, student_file_path: str) -> Dict[str, Any]:
        """Score a single student assignment."""
        print(f"Processing: {student_file_path}")

        # Extract text from document
        student_text = self.extract_text_from_docx(student_file_path)

        # Check if submission is complete
        is_complete, reason = self.check_if_submission_is_complete(student_text)

        if not is_complete:
            # Return a very low score for incomplete submissions
            print(f"  ⚠️  Incomplete submission detected: {reason}")
            evaluation = {
                "total_score": "1 / 25",
                "category_breakdown": {
                    "prompt_clarity": "0/5",
                    "real_world_relevance": "0/5",
                    "reflection_quality": "0/5",
                    "responsible_use": "0/5",
                    "writing_clarity": "1/5"
                },
                "summary_feedback": f"This submission appears to be incomplete. {reason} To receive a proper evaluation, please submit a complete assignment that includes all three sections with your own work: Section 1 (Understanding Generative AI), Section 2 (Writing and Testing Prompts with Steps 1-5), and Section 3 (Responsible and Safe Use).",
                "improvement_suggestion": "Please review the assignment requirements and ensure you complete all sections with thoughtful, original responses. Focus on providing real examples from your work and detailed reflections on your learning.",
                "raw_response": f"Incomplete submission: {reason}"
            }
        else:
            # Get evaluation from AI
            evaluation = self.evaluate_assignment(student_text)

        # Add metadata
        evaluation["metadata"] = {
            "student_file": os.path.basename(student_file_path),
            "evaluation_date": datetime.now().isoformat()
        }

        return evaluation

    def score_all_assignments(self, input_dir: str = "student-assignments",
                            output_dir: str = "evaluation-results") -> List[Dict[str, Any]]:
        """Score all assignments in the input directory."""
        # Create output directory if it doesn't exist
        Path(output_dir).mkdir(exist_ok=True)

        # Find all .docx files
        input_path = Path(input_dir)
        docx_files = list(input_path.glob("*.docx"))

        if not docx_files:
            print(f"No .docx files found in {input_dir}")
            return []

        print(f"Found {len(docx_files)} assignment(s) to evaluate.")

        results = []

        for docx_file in docx_files:
            try:
                # Score the assignment
                evaluation = self.score_assignment(str(docx_file))
                results.append(evaluation)

                # Save individual result
                student_name = docx_file.stem
                output_file = Path(output_dir) / f"{student_name}_evaluation.json"

                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(evaluation, f, indent=2, ensure_ascii=False)

                print(f"✓ Saved evaluation to: {output_file}")
                print(f"  Score: {evaluation['total_score']}")
                print()

            except Exception as e:
                print(f"✗ Error processing {docx_file}: {str(e)}")
                print()

        # Save summary of all results
        summary_file = Path(output_dir) / "all_evaluations_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)

        print(f"✓ All evaluations saved to: {output_dir}")
        print(f"✓ Summary saved to: {summary_file}")

        return results


def main():
    """Main function to run the automated scoring."""
    print("=" * 60)
    print("Automated Assignment Scoring System")
    print("=" * 60)
    print()

    scorer = AssignmentScorer()
    results = scorer.score_all_assignments()

    print()
    print(f"Completed! Evaluated {len(results)} assignment(s).")


if __name__ == "__main__":
    main()
